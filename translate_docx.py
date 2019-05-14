# coding:utf-8
from googletrans import Translator
import requests
import warnings
from threading import Thread
import re
from queue import Queue
import docx
import pickle
import time
import os

warnings.filterwarnings('ignore')


class SahajaTrans(Thread):
    def __init__(self, queue, dic):
        Thread.__init__(self)
        self.queue = queue
        self.dic = dic

    def run(self):
        while self.queue.qsize():
            print('Queue size: ', self.queue.qsize())
            idx = self.queue.get()
            self.google_translate(idx)
            self.queue.task_done()

    def google_translate(self, idx):
        try:
            content = self.dic[idx]['eng']
            raw_list = split_to_limit_list(content)
            raw_list_num = len(raw_list)
            max_list_num = 400
            trans_txt = ''

            # 判断如果列表太长，就按每400个一个代理进行处理，以免服务器拒绝访问
            for i in range(raw_list_num):
                part_list = raw_list[i * max_list_num: (i + 1) * max_list_num]
                proxy = get_proxy()
                proxies = {"http": "http://{}".format(proxy), "https": "https://{}".format(proxy)}

                translator = Translator(service_urls=['translate.google.cn'], proxies=proxies, timeout=30)
                translations = translator.translate(part_list, dest='zh-cn')
                tran_list = []
                # 判断如果成员不是字符，而是列表的说明不是按行分割的，则按所有原文+所有译文来展示，由于译文没有换行符，手工增加
                for translation in translations:
                    if isinstance(translation, list):
                        origin_list, text_list = [], []
                        for tra in translation:
                            origin_list.append(tra.origin)
                            text_list.append(tra.text)
                            if '\n' in tra.origin:
                                text_list.append('\n\n')
                        tran_list.extend(origin_list)
                        tran_list.extend('\r')
                        tran_list.extend(text_list)
                    else:
                        tran_list.append(translation.origin)
                        tran_list.extend('\r')
                        tran_list.append(translation.text)
                        if '\n' in translation.origin:
                            tran_list.append('\n\n')
                trans_txt += ''.join(tran_list)
            self.dic[idx]['trans'] = trans_txt
            print(idx, 'translation completed.')
        except Exception as e:
            print(idx, e)
            self.queue.put(idx)


def get_proxy():
    return requests.get("http://127.0.0.1:5010/get/").text


# 把字符按换行符拆分成列表，如果其中成员长度超过5000就拆分成小于5000的成员，如[100, [3000, 2000], 2000]
def split_to_limit_list(content):
    limit_length = 5000
    lines = content.splitlines(keepends=True)
    part_num = len(lines)
    for i in range(part_num):
        if len(lines[i]) > limit_length:
            sentences = iter(re.split(r'([.])', lines[i]))
            sen_list, current = [], next(sentences)
            for sentence in sentences:
                if len(current) + len(sentence) >= limit_length:
                    sen_list.append(current)
                    current = sentence
                else:
                    current += sentence
            sen_list.append(current)
            lines[i] = sen_list
    return lines


def read_to_dic(docu):
    dic = {}
    for i in range(len(docu.paragraphs)):
        if docu.paragraphs[i].text:
            print(len(dic))
            dic[i] = {'eng': docu.paragraphs[i].text}
    return dic


def dic_to_file(docu, save_name, dic):
    print('Start save to docx.')
    for i in range(len(docu.paragraphs)):
        if i in dic:
            docu.paragraphs[i].text = dic[i].get('trans', '')
    docu.save(save_name)


def check_complete(dic):
    for d in dic.values():
        if 'trans' not in d:
            return False
    else:
        return True


def main():
    pk_file = 'dic.pickle'
    file_name = '女神往世书全本 SrimadDeviPurana.docx'
    save_name = '女神往世书全本 SrimadDeviPurana_trans.docx'

    docu = docx.Document(file_name)

    if os.path.exists(pk_file):
        with open(pk_file, 'rb') as f:
            dic = pickle.load(f)
    else:
        dic = read_to_dic(docu)

    queue = Queue()

    for i in dic:
        if 'trans' not in dic[i]:
            queue.put(i)

    for _ in range(30):
        sa = SahajaTrans(queue, dic)
        sa.start()

    while not check_complete(dic):
        time.sleep(60)
        with open(pk_file, 'wb') as f:
            pickle.dump(dic, f)

    # queue.join()
    dic_to_file(docu, save_name, dic)
    print('Done')


if __name__ == '__main__':
    main()
